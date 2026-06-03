#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成物流管理系统实验报告 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ========== 全局样式设置 ==========
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, indent=False, font_size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    return p

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

# ========== 封面 ==========
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('面向对象程序设计实践（C++）')
run.font.size = Pt(26)
run.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('课程实验报告')
run.font.size = Pt(22)
run.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

for _ in range(4):
    doc.add_paragraph()

# 信息表
table = doc.add_table(rows=4, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
data = [
    ('题  目', '物流管理平台设计与实现'),
    ('姓  名', '（请填写）'),
    ('学  号', '（请填写）'),
    ('班  级', '（请填写）'),
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v
    for cell in table.rows[i].cells:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(14)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ========== 目录页 ==========
add_heading_styled('目录', level=1)
toc_items = [
    ('1  前言', 3),
    ('  1.1 题目介绍', 3),
    ('  1.2 需求理解', 3),
    ('2  总体设计', 4),
    ('  2.1 子系统划分', 4),
    ('  2.2 子系统关系', 4),
    ('3  详细设计', 5),
    ('  3.1 题目一子系统详细设计', 5),
    ('  3.2 题目二子系统详细设计', 5),
    ('  3.3 题目三子系统详细设计', 5),
    ('  3.4 数据库说明', 6),
    ('  3.5 接口协议说明', 6),
    ('4  实现', 7),
    ('  4.1 主要问题和解决方案', 7),
    ('  4.2 想法', 7),
    ('  4.3 经验', 7),
    ('  4.4 教训', 7),
]
for text, page in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# ========== 1 前言 ==========
add_heading_styled('1  前言', level=1)

add_heading_styled('1.1 题目介绍', level=2)
add_para('本实验的任务是使用C++语言，基于面向对象的程序设计方法，设计并实现一个简单的物流管理平台，提供物流管理、用户管理、员工管理等功能。整个作业包含三个递进式题目，使得物流管理平台功能逐步增强：', indent=True)
add_para('• 题目一：物流业务管理子系统和用户管理子系统（封装）', indent=True)
add_para('   功能包括用户注册登录、修改密码、余额管理、发送快递（固定15元）、接收快递、查询快递、管理员物流管理。', indent=True)
add_para('• 题目二：快递员任务管理子系统（继承与多态）', indent=True)
add_para('   引入快递员角色和快递分类体系（易碎品8元/kg、图书2元/本、普通5元/kg），管理员分配揽收任务，快递员揽收后获得50%分成。', indent=True)
add_para('• 题目三：物流管理系统网络版（C/S架构与并发）', indent=True)
add_para('   将单机版改造为Client-Server架构，使用Winsock TCP通信，服务端多线程并发处理多客户端请求。', indent=True)

add_heading_styled('1.2 需求理解', level=2)
add_para('本作业的核心要求可以分为三个维度：', indent=True)
add_para('（1）面向对象设计：充分利用封装、继承、多态三大特性。题目一体现封装——将数据和操作封装在类中；题目二体现继承与多态——通过Person基类消除重复代码，通过Package抽象基类和纯虚函数实现多态计费；题目三体现复杂架构下的面向对象设计——将业务逻辑、数据管理、网络通信分离为独立类。', indent=True)
add_para('（2）数据持久化：使用文件存储用户、快递、管理员、快递员等信息。每个题目的数据文件格式逐步演进，体现系统的迭代开发过程。', indent=True)
add_para('（3）C/S架构：要求使用传统的Client-Server结构，基于Socket通信而非RPC框架，体现对网络编程底层原理的理解。', indent=True)
add_para('（4）代码规范：所有不改变对象状态的成员函数均标注const，不允许出现非类成员函数（除主函数和必要的友元函数外）。', indent=True)

doc.add_page_break()

# ========== 2 总体设计 ==========
add_heading_styled('2  总体设计', level=1)

add_heading_styled('2.1 子系统划分', level=2)
add_para('本系统按照三个题目划分为三个子系统，每个子系统在上一题基础上进行功能增强和架构升级：', indent=True)

# 总体设计表
t = doc.add_table(rows=4, cols=5)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['子系统', '核心概念', '架构', '角色数量', '快递状态']
for i, h in enumerate(headers):
    cell = t.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True

data_rows = [
    ['题目一', '封装', '单机单体', '2（用户/管理员）', '2种（待签收→已签收）'],
    ['题目二', '继承+多态', '单机MVC', '3（+快递员）', '3种（待揽收→待签收→已签收）'],
    ['题目三', '网络+并发', 'Client-Server', '3（同题目二）', '3种（同题目二）'],
]
for ri, row_data in enumerate(data_rows):
    for ci, val in enumerate(row_data):
        t.rows[ri+1].cells[ci].text = val
        for p in t.rows[ri+1].cells[ci].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading_styled('2.2 子系统关系', level=2)
add_para('三个子系统呈现递进关系：', indent=True)
add_para(' 题目一（封装） ──→ 题目二（继承+多态） ──→ 题目三（C/S架构+并发）', bold=True)
add_para('  单体应用              MVC分层              网络分布式', indent=True)
add_para('   固定计费             多态计费              多客户端并发', indent=True)
add_para('   2种角色              3种角色               3种角色保持不变', indent=True)
add_para('   2种状态              3种状态               3种状态保持不变', indent=True)

add_para('')
add_para('题目三的最终系统架构图：', bold=True)
add_code('''
┌──────────┐      TCP:8888       ┌──────────────────────────┐
│ Client A  │───────→─────────────│         Server           │
│ (用户界面)  │←───────────────────│  ├─ DataManager          │→ users.txt
├──────────┤                     │  │  (唯一数据源)           │→ packages.txt
│ Client B  │───────→─────────────│  ├─ CreateThread线程池    │→ couriers.txt
│ (管理员)   │←───────────────────│  ├─ CRITICAL_SECTION锁    │→ admin.txt
├──────────┤                     │  └─ Handlers 请求分发      │→ next_*.txt
│ Client C  │───────→─────────────│                           │
│ (快递员)   │←───────────────────│ 业务逻辑全部在服务端       │
└──────────┘                     └──────────────────────────┘
''')

add_para('各子系统的类体系演进：', bold=True)
add_code('''
题目一                            题目二                           题目三
User (独立) ────→ User : Person ────→ 同题目二，网络传输
                    Person (新增)
Admin (独立) ────→ Admin : Person ──→ 同题目二
                    Courier : Person ──→ 同题目二(新增)
Package (具体类) ──→ Package (抽象基类) ───→ 同题目二
                      ├── FragilePackage
                      ├── BookPackage
                      └── NormalPackage
''')

doc.add_page_break()

# ========== 3 详细设计 ==========
add_heading_styled('3  详细设计', level=1)

# -------- 3.1 题目一 --------
add_heading_styled('3.1 题目一子系统详细设计', level=2)

add_para('3.1.1 类图', bold=True)
add_code('''
┌──────────────┐   ┌───────────────┐   ┌────────────────────┐
│    User      │   │    Admin      │   │      Package       │
├──────────────┤   ├───────────────┤   ├────────────────────┤
│- username_   │   │- username_    │   │- packageId_        │
│- name_       │   │- name_        │   │- sender_           │
│- phonenum_   │   │- password_    │   │- receiver_         │
│- password_   │   │- balance_     │   │- sendTime_         │
│- balance_    │   ├───────────────┤   │- receiveTime_      │
│- address_    │   │+ CheckPwd()   │   │- status_(0/1)      │
├──────────────┤   │+ SetPwd()     │   │- description_      │
│+ CheckPwd()  │   │+ AddBalance() │   ├────────────────────┤
│+ Recharge()  │   └───────────────┘   │+ Sign()            │
│+ Deduct()    │         |              └────────────────────┘
└──────┬───────┘         |
       |                 |
       └─── LogisticsSystem (外观类，含所有业务逻辑+UI+数据持久化)
''')

add_para('3.1.2 关键类的设计', bold=True)
add_para('User类：表示平台用户，包含用户名、姓名、电话、密码、余额、地址等属性。提供余额查询/充值/扣费、密码验证/修改等接口。所有getter方法均标注const，不修改对象状态。', indent=True)
add_para('Admin类：表示物流公司管理员，包含用户名、姓名、密码、公司余额等属性。无需注册功能，预设admin/admin123。', indent=True)
add_para('Package类：表示快递包裹，包含单号、寄件人、收件人、寄送时间、接收时间、状态（0待签收/1已签收）、物品描述等属性。提供Sign()方法将状态从未签收改为已签收。', indent=True)
add_para('LogisticsSystem类：系统主控类，约700行代码。负责主菜单循环、业务逻辑调度、数据持久化读写。缺点在于职责过于集中——同时承担UI、业务逻辑、数据管理三个角色。', indent=True)

add_para('3.1.3 业务流程', bold=True)
add_para('用户注册→登录→发送快递（固定15元，自动扣费）→快递状态=待签收→收件人查看未签收快递→选择签收→状态变为已签收。管理员登录后可查看所有用户信息和所有快递信息。', indent=True)

add_para('3.1.4 数据文件格式', bold=True)
add_code('''
users.txt:     username|name|phone|password|balance|address
admin.txt:     username|name|password|balance
packages.txt:  id|sender|receiver|sendTime|receiveTime|status|description
next_id.txt:   N （自增ID计数器）
''')

# -------- 3.2 题目二 --------
add_heading_styled('3.2 题目二子系统详细设计', level=2)

add_para('3.2.1 类继承体系', bold=True)
add_code('''
                    ┌──────────────────┐
                    │     Person       │  ← 抽象基类
                    ├──────────────────┤
                    │- username_       │
                    │- name_           │
                    │- password_       │
                    │- balance_        │
                    └───────┬──────────┘
                            │
          ┌─────────────────┼─────────────────┐
          ▼                 ▼                  ▼
   ┌─────────────┐   ┌─────────────┐   ┌─────────────┐
   │    User     │   │    Admin    │   │   Courier   │
   ├─────────────┤   ├─────────────┤   ├─────────────┤
   │- phonenum_ │   │(无额外属性)  │   │- id_        │
   │- address_  │   └─────────────┘   │- phone_     │
   └─────────────┘                    └─────────────┘

                    ┌─────────────────────┐
                    │   Package(abstract) │  ← 抽象基类
                    ├─────────────────────┤
                    │+ GetPrice() = 0     │  ← 纯虚函数
                    └─────────┬───────────┘
                              │
        ┌─────────────────────┼─────────────────────┐
        ▼                     ▼                      ▼
┌───────────────┐   ┌───────────────┐   ┌──────────────┐
│FragilePackage │   │  BookPackage  │   │NormalPackage │
├───────────────┤   ├───────────────┤   ├──────────────┤
│- weight_      │   │- count_       │   │- weight_     │
│GetPrice():    │   │GetPrice():    │   │GetPrice():   │
│  8.0 × weight │   │  2.0 × count │   │  5.0 × weight│
└───────────────┘   └───────────────┘   └──────────────┘
''')

add_para('3.2.2 MVC分层架构', bold=True)
add_code('''
┌────────────────────────────────────┐
│     LogisticsSystem (外观)         │
│  持有 DataManager + 3 Controller   │
└──────────────┬─────────────────────┘
               │
   ┌───────────┼───────────┐
   ▼           ▼           ▼
┌──────┐ ┌──────┐ ┌──────┐
│User  │ │Admin │ │Courier│  ← 控制器层（菜单+业务逻辑）
│Ctrl  │ │Ctrl  │ │Ctrl   │
└──┬───┘ └──┬───┘ └──┬────┘
   │        │        │
   └────────┼────────┘
            ▼
    ┌──────────────┐
    │ DataManager  │  ← 数据层（文件读写+集合管理）
    ├──────────────┤
    │ vector<Package*>   │ ← 多态存储
    │ vector<User>       │
    │ vector<Courier>    │
    └──────┬───────┘
''')

add_para('3.2.3 关键类的设计', bold=True)
add_para('Person类（基类）：抽取User、Admin、Courier的共同属性和方法（username_, name_, password_, balance_及其getter/setter）。使用protected访问权限，允许子类直接访问基类成员。', indent=True)
add_para('Package类（抽象基类）：将Package改为抽象类，添加纯虚函数virtual double GetPrice() const = 0。新增courierId_字段关联揽收快递员。新增IsWaitingCollect()判断方法。状态增加为3种：0=待揽收、1=待签收、2=已签收。', indent=True)
add_para('FragilePackage/BookPackage/NormalPackage：分别实现GetPrice()虚函数，计算公式为8.0*weight、2.0*count、5.0*weight。体现了多态的应用——LogisticsSystem通过基类指针调用GetPrice()，实际执行的是子类版本。', indent=True)
add_para('Courier类：继承Person，增加id_（快递员编号）和phone_（电话）属性。快递员使用ID+密码登录。', indent=True)
add_para('DataManager类：单一职责原则——仅负责数据管理。封装所有文件读写操作，管理users_、couriers_、packages_（vector<Package*>多态存储）、admin_等集合。提供findUser、findCourier、findPackage等查询方法。', indent=True)
add_para('UserController/AdminController/CourierController：分别对应三种角色的业务逻辑，包含菜单循环和操作调度。通过引用持有DataManager，实现数据访问。', indent=True)

add_para('3.2.4 快递状态机', bold=True)
add_code('''
待揽收(0) ──(快递员揽收)──→ 待签收(1) ──(用户签收)──→ 已签收(2)
''')

add_para('3.2.5 揽收业务流程', bold=True)
add_para('（1）用户选择快递类型→输入收件人、重量/数量→系统调用对应子类的GetPrice()计算费用→扣费→状态=待揽收。', indent=True)
add_para('（2）管理员查看待揽收快递列表→选择快递→选择快递员→分配（记录courierId，状态不变）。', indent=True)
add_para('（3）快递员登录→查看我的任务（courierId匹配的待揽收快递）→选择揽收→状态变为待签收→揽收费=快递费×50%→管理员扣费→快递员入账。', indent=True)

add_para('3.2.6 数据文件格式（相对于题目一的变化）', bold=True)
add_code('''
packages.txt:  id|sender|receiver|sendTime|receiveTime|status|type|detail|description|courierId
              // type: 1=易碎品 2=图书 3=普通
              // detail: 重量(kg)或数量(本)
couriers.txt:  id|name|phone|password|balance  (新增)
next_courier_id.txt:  N  (新增)
''')

# -------- 3.3 题目三 --------
add_heading_styled('3.3 题目三子系统详细设计', level=2)

add_para('3.3.1 架构变更：从单机到C/S', bold=True)
add_code('''
题目二（单机）                    题目三（网络版）

┌─────────────────┐         ┌──────────┐    TCP:8888    ┌──────────┐
│ LogisticsSystem │         │ Client A │─────→────│ Server   │
│  + DataManager  │         └──────────┘          ├──────────┤
│  + Controllers  │         ┌──────────┐          │DataManager│
│  + Console UI   │         │ Client B │─────→────│+线程池    │
└─────────────────┘         └──────────┘          │+锁       │
                              ┌──────────┐          └──────────┘
 业务逻辑+UI+数据全部耦合      │ Client C │─────→────
                              └──────────┘
                              客户端是"哑终端"      服务端处理所有业务逻辑
''')

add_para('3.3.2 关键类的设计', bold=True)
add_para('Server类：服务端主类，封装Winsock生命周期（WSAStartup→socket→bind→listen→accept循环）。每接受一个客户端连接，创建一个新线程（CreateThread）处理该客户端的请求。使用CRITICAL_SECTION保护DataManager的并发访问。', indent=True)
add_para('Client类：客户端主类，封装连接建立（connect）、请求发送（sendRequest）、响应接收（receiveResponse）。不包含任何业务逻辑，所有操作通过协议字符串委托给服务端。提供自动重连机制（reconnect）。', indent=True)
add_para('Protocol命名空间：定义通信协议常量（30个命令名、分隔符、响应前缀）。提供buildRequest/buildOkResponse/buildErrResponse/buildDataResponse辅助函数，统一消息格式。', indent=True)
add_para('Handler函数集群：按角色拆分为UserHandlers.cpp（13个handler）、AdminHandlers.cpp（13个handler）、CourierHandlers.cpp（5个handler）。每个handler接收字符串参数，返回字符串响应，不涉及任何UI操作。', indent=True)

add_para('3.3.3 通信协议', bold=True)
add_code('''
请求格式:  COMMAND|param1|param2|...\\n
成功响应:  OK|message\\n[data\\n]END\\n
错误响应:  ERR|message\\n

示例:
客户端→服务端: LOGIN|user|zhangsan|123456\\n
服务端→客户端: OK|用户登录成功，欢迎张三\\n

客户端→服务端: SEND_PACKAGE|lisi|1|2.5|衣物\\n
服务端→客户端: OK|包裹已发送！ID:3，费用:20.0元，余额:80.0元\\n

两阶段协议（揽收/签收）:
阶段1: 客户端→RECEIVE_PACKAGE\\n
        服务端→OK|待签收包裹\\n0|1|zhangsan|衣物\\nEND\\n
阶段2: 客户端→RECEIVE_PACKAGE|0\\n
        服务端→OK|已签收1个包裹\\n
''')

add_para('3.3.4 服务端多线程设计', bold=True)
add_code('''
主线程: accept() 循环
  │
  ├─ 客户端A连接 → CreateThread → handleClient(A)
  │                                  ├─ while(recv → processRequest → send)
  │                                  │      │
  │                                  │      ├─ EnterCriticalSection(&cs)
  │                                  │      ├─ 按userType分发:
  │                                  │      │   0(未登录)→dispatchUnauthCommand
  │                                  │      │   1(用户)  →dispatchUserCommand
  │                                  │      │   2(管理员)→dispatchAdminCommand
  │                                  │      │   3(快递员)→dispatchCourierCommand
  │                                  │      ├─ LeaveCriticalSection(&cs)
  │                                  │      └─ 若需保存→单独加锁saveData()
  │                                  └─ 断开→closesocket→线程结束
  │
  ├─ 客户端B连接 → CreateThread → handleClient(B)...
  └─ 客户端C连接 → CreateThread → handleClient(C)...
''')

# -------- 3.4 数据库说明 --------
add_heading_styled('3.4 数据库说明', level=2)

add_para('本系统使用文本文件进行数据持久化，不使用关系数据库。以下为所有数据文件的格式说明：', indent=True)

t = doc.add_table(rows=7, cols=4)
t.style = 'Table Grid'
headers2 = ['文件名', '字段格式', '说明', '所属题目']
for i, h in enumerate(headers2):
    t.rows[0].cells[i].text = h
    for p in t.rows[0].cells[i].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs: r.bold = True

file_data = [
    ['users.txt', 'username|name|phone|password|balance|address', '用户信息', '题目一/二/三'],
    ['admin.txt', 'username|name|password|balance', '管理员信息', '题目一/二/三'],
    ['packages.txt(题目一)', 'id|sender|receiver|sendTime|receiveTime|status|desc', '快递(2状态)', '题目一'],
    ['packages.txt(题目二/三)', 'id|sender|receiver|sendTime|receiveTime|status|type|detail|desc|courierId', '快递(3状态+类型)', '题目二/三'],
    ['couriers.txt', 'id|name|phone|password|balance', '快递员信息', '题目二/三'],
    ['next_id.txt / next_courier_id.txt', 'N', '自增ID计数器', '全部'],
]
for ri, row_data in enumerate(file_data):
    for ci, val in enumerate(row_data):
        t.rows[ri+1].cells[ci].text = val
        for p in t.rows[ri+1].cells[ci].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

# -------- 3.5 接口协议说明 --------
add_heading_styled('3.5 接口协议说明', level=2)

add_para('题目三中，客户端与服务端之间使用自定义的纯文本协议进行通信。', indent=True)

add_para('3.5.1 协议语法', bold=True)
add_code('''
• 分隔符：'|'（竖线）
• 终止符：'\\n'（换行）
• 请求格式：COMMAND|param1|param2|...\\n
• 成功响应：OK|message\\n
• 数据响应：OK|message\\n<data>\\nEND\\n
• 错误响应：ERR|message\\n
''')

add_para('3.5.2 协议语义（命令全表）', bold=True)

t2 = doc.add_table(rows=9, cols=3)
t2.style = 'Table Grid'
for i, h in enumerate(['分类', '命令', '功能']):
    t2.rows[0].cells[i].text = h
    for p in t2.rows[0].cells[i].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs: r.bold = True

cmd_data = [
    ['通用', 'LOGIN / REGISTER / LOGOUT / EXIT', '登录/注册/退出/断开'],
    ['用户', 'SEND_PACKAGE / RECEIVE_PACKAGE', '发送/签收快递'],
    ['用户', 'QUERY_SENT / QUERY_RECEIVED', '查询发出/收到的快递'],
    ['用户', 'QUERY_BY_ID / QUERY_BY_TIME_*', '按单号/时间查询'],
    ['用户', 'GET_BALANCE / RECHARGE / CHANGE_PWD', '余额查询/充值/改密'],
    ['管理员', 'ADMIN_SHOW_USERS / ADMIN_SHOW_PACKAGES', '查看用户/快递'],
    ['管理员', 'ADMIN_ADD/REMOVE/SHOW_COURIER', '快递员管理'],
    ['快递员', 'COURIER_MY_TASKS / COLLECT / MY_RECORDS', '任务/揽收/记录'],
]
for ri, row_data in enumerate(cmd_data):
    for ci, val in enumerate(row_data):
        t2.rows[ri+1].cells[ci].text = val

add_para('')
add_para('3.5.3 底层承载协议', bold=True)
add_para('协议承载于TCP（传输控制协议）之上，使用Winsock2 API实现。服务端监听8888端口，客户端通过IPv4连接。TCP提供面向连接、可靠传输的字节流服务。', indent=True)

doc.add_page_break()

# ========== 4 实现 ==========
add_heading_styled('4  实现', level=1)

add_heading_styled('4.1 实现过程中遇到的主要问题和解决方案', level=2)

add_para('问题1：代码重复（题目一→题目二的改进）', bold=True)
add_para('题目一中User和Admin的username_、name_、password_、balance_属性完全重复，各自实现相同的getter/setter。解决方案是在题目二中引入Person基类，将公共属性和方法提取到基类中，User和Admin通过继承复用代码。这个重构体现了封装和继承的合理运用。', indent=True)

add_para('问题2：计费逻辑的扩展性（题目一→题目二的改进）', bold=True)
add_para('题目一的计费逻辑是硬编码的固定15元，新增快递类型需要修改LogisticsSystem类的sendPackage方法。解决方案是将Package改为抽象基类，定义纯虚函数GetPrice() = 0，每种快递类型通过子类实现自己的计费逻辑。新增快递类型时只需新增子类，无需修改现有代码——体现了开闭原则。', indent=True)

add_para('问题3：非法输入处理', bold=True)
add_para('C++的cin在输入非数字字符时会导致输入流进入错误状态，后续所有输入操作都会失败。解决方案是每次cin >> num后检查failbit，如果出错则调用cin.clear()清除错误标志，cin.ignore()丢弃缓冲区中的错误输入，然后提示用户重新输入。对于菜单选择，将所有choice改为string类型，避免类型转换错误。', indent=True)

add_para('问题4：多线程数据竞争（题目三）', bold=True)
add_para('服务端多个客户端线程同时访问DataManager时，可能发生数据竞争（如两个线程同时修改同一个用户的余额）。解决方案是使用Windows CRITICAL_SECTION（临界区）保护所有对共享数据的操作。每次请求处理在EnterCriticalSection/LeaveCriticalSection之间执行，确保同一时刻只有一个线程在操作数据。数据写文件操作在临界区外执行，避免文件I/O阻塞其他客户端的请求处理。', indent=True)

add_para('问题5：断线重连（题目三）', bold=True)
add_para('客户端在运行过程中可能因网络问题断开连接（如服务端重启）。解决方案是实现自动重连机制：当recv失败时，客户端自动尝试重新连接，最多重试5次，每次间隔2秒。重连成功后自动恢复操作，无需用户手动重启客户端。', indent=True)

add_para('问题6：端口占用（题目三）', bold=True)
add_para('服务端关闭后重新启动时，操作系统可能仍处于TIME_WAIT状态，导致bind失败。解决方案是在socket创建后设置SO_REUSEADDR选项，允许端口立即重用。', indent=True)

add_heading_styled('4.2 想法', level=2)
add_para('（1）面向对象设计不是一蹴而就的。从题目一到题目三的演进过程本身就是最好的教学——开始的简单设计必然存在不足，随着需求增加逐步重构优化，这是实际软件开发中的常态。', indent=True)
add_para('（2）纯文本协议的优势。在网络版设计中，选择纯文本协议而非JSON或二进制协议，虽然看起来"原始"，但调试极为方便——可以直接用telnet或Python socket发送命令测试，不需要额外的序列化/反序列化库。', indent=True)
add_para('（3）const正确性的重要性。在重构过程中，我们发现一个函数是否标注const会影响整个调用链——如果调用一个非const函数，调用者也必须是非const的。设计初期就考虑const正确性，可以避免后期大量的接口修改。', indent=True)

add_heading_styled('4.3 经验', level=2)
add_para('（1）先设计再编码。在开始写代码之前，画出类图、理清继承关系、设计好接口，可以大幅减少后期的返工。题目一因为缺乏前期设计，LogisticsSystem类过于臃肿（约700行），题目二重构时将其拆分成了多个职责明确的类。', indent=True)
add_para('（2）分层架构的价值。将系统分为数据层（DataManager）、业务层（Controller）、表现层（菜单UI），每层各司其职，修改其中一层不需要改动其他层。题目三将表现层彻底分离到客户端，进一步印证了分层设计的优势。', indent=True)
add_para('（3）自动化测试。题目三的Python自动化测试脚本（test_demo.py和concurrency_test.py）可以快速验证系统功能，比手动输入测试高效得多。特别是并发测试，手动模拟10个客户端同时操作几乎不可能，脚本可以在几秒内完成。', indent=True)

add_heading_styled('4.4 教训', level=2)
add_para('（1）不要在初期过度设计。题目一直接使用简单类设计，虽然导致了题目二的大量重构，但这是合理的——如果一开始就引入抽象基类和MVC分层，反而会让初学者难以理解封装的基本概念。三个题目的递进式设计正好符合学习曲线。', indent=True)
add_para('（2）文件格式兼容性需要注意。题目一和题目二的packages.txt格式不同（字段数不同），如果混用会导致加载错误。在实际开发中应设计向前兼容的文件格式，或在版本升级时提供数据迁移脚本。', indent=True)
add_para('（3）异常安全。在题目三的handler中，如果new Package失败或std::stoi转换抛异常，可能会导致死锁（已加锁但未解锁）或内存泄漏。解决方案是使用try-catch包裹整个处理流程，在catch块中确保解锁并清理资源。', indent=True)

# ========== 附录 ==========
doc.add_page_break()
add_heading_styled('附录：程序编译与运行说明', level=1)

add_para('一、环境要求', bold=True)
add_para('• 编译器：MinGW-w64 (g++ 8.0+)，支持C++11', indent=True)
add_para('• 系统：Windows 10/11', indent=True)
add_para('• 库依赖：ws2_32（Winsock2，仅题目三需要）', indent=True)

add_para('二、题目一编译运行', bold=True)
add_code('cd topic1')
add_code('g++ -std=c++11 -Wall -o topic1.exe src/main.cpp src/*.cpp -I./include')
add_code('topic1.exe')

add_para('三、题目二编译运行', bold=True)
add_code('cd topic2')
add_code('g++ -std=c++11 -Wall -o topic2.exe src/main.cpp src/*.cpp -I./include')
add_code('topic2.exe')

add_para('四、题目三编译运行', bold=True)
add_code('cd topic3')
add_code('REM 编译服务端')
add_code('g++ -std=c++11 -Wall -o server/server.exe server/main.cpp src/Server.cpp src/UserHandlers.cpp src/AdminHandlers.cpp src/CourierHandlers.cpp src/Person.cpp src/User.cpp src/Admin.cpp src/Courier.cpp src/Package.cpp src/NormalPackage.cpp src/FragilePackage.cpp src/BookPackage.cpp src/DataManager.cpp -I./include -lws2_32')
add_code('')
add_code('REM 编译客户端')
add_code('g++ -std=c++11 -Wall -o client/client.exe client/main.cpp src/Client.cpp src/Person.cpp src/Package.cpp -I./include -lws2_32')
add_code('')
add_code('REM 运行')
add_code('start server/server.exe')
add_code('start client/client.exe   # 可启动多个客户端')

# ========== 保存文件 ==========
output_path = r'C:\Code\oop-logistics-system\docs\实验报告.docx'
doc.save(output_path)
print(f'实验报告已生成: {output_path}')
